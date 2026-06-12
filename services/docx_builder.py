"""
services/docx_builder.py — Renders structured JSON into ATS-friendly DOCX files.

ATS rules enforced:
  - Single column, no layout tables, no text boxes
  - Native Word list styles for bullets (never manual • characters)
  - Important content never in headers/footers
  - Standard fonts (Calibri) — universally readable by parsers
  - Contact info as plain body text, not in document header

Design language: dark navy (#1E3A5F) headings, gold (#C9A84C) rule under name,
light grey section dividers — matches candidate's existing brand.
"""
import io
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Brand colours ───────────────────────────────────────────
NAVY   = RGBColor(0x1E, 0x3A, 0x5F)
GOLD   = RGBColor(0xC9, 0xA8, 0x4C)
BLACK  = RGBColor(0x1A, 0x1A, 0x1A)
GREY   = RGBColor(0x55, 0x55, 0x55)
LGREY  = RGBColor(0x88, 0x88, 0x88)

def _set_font(run, size_pt, bold=False, italic=False, color=None):
    run.font.name   = "Calibri"
    run.font.size   = Pt(size_pt)
    run.font.bold   = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color

def _para_spacing(para, before=0, after=0, line=None):
    fmt = para.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after  = Pt(after)
    if line:
        from docx.shared import Pt as _Pt
        fmt.line_spacing = _Pt(line)

def _add_bottom_border(para, color_hex="C9A84C", size=12):
    """Add a bottom border to a paragraph (used under name)."""
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)

def _add_section_rule(para):
    """Thin navy line under section header."""
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1E3A5F")
    pBdr.append(bottom)
    pPr.append(pBdr)

def _add_tab_stop(para, position_inches, alignment="right"):
    """Right-aligned tab stop for dates."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches
    pPr = para._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), alignment)
    tab.set(qn("w:pos"), str(int(position_inches * 1440)))
    tabs.append(tab)
    pPr.append(tabs)

def _set_margins(doc, top=1.0, bottom=1.0, left=1.0, right=1.0):
    for section in doc.sections:
        section.top_margin    = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin   = Inches(left)
        section.right_margin  = Inches(right)

def _section_header(doc, text):
    p = doc.add_paragraph()
    _para_spacing(p, before=10, after=2)
    _add_section_rule(p)
    run = p.add_run(text.upper())
    _set_font(run, 10.5, bold=True, color=NAVY)
    return p

def _bullet_para(doc, text):
    """ATS-safe bullet: native Word list style, no manual characters."""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.left_indent  = Inches(0.25)
    run = p.add_run(text)
    _set_font(run, 10.5, color=BLACK)
    return p

# ── CV builder ───────────────────────────────────────────────
def build_cv(data: dict) -> bytes:
    """
    data: structured CV JSON from the AI.
    Returns DOCX file bytes ready for HTTP response.
    """
    doc = Document()
    _set_margins(doc, top=0.75, bottom=0.75, left=0.85, right=0.85)

    # Remove default empty paragraph
    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)

    # ── NAME ──────────────────────────────────────────────────
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_spacing(name_para, before=0, after=4)
    _add_bottom_border(name_para, color_hex="C9A84C", size=16)
    name_run = name_para.add_run(data.get("name", "").upper())
    _set_font(name_run, 18, bold=True, color=NAVY)

    # ── CONTACT LINE ──────────────────────────────────────────
    contact_parts = []
    if data.get("phone"):
        contact_parts.append(data["phone"])
    if data.get("email"):
        contact_parts.append(data["email"])
    if data.get("linkedin"):
        contact_parts.append(data["linkedin"])
    if data.get("location"):
        contact_parts.append(data["location"])
    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_spacing(contact_para, before=3, after=6)
    contact_run = contact_para.add_run("  |  ".join(contact_parts))
    _set_font(contact_run, 9.5, color=GREY)

    # ── PROFESSIONAL SUMMARY ──────────────────────────────────
    summary = (data.get("summary") or "").strip()
    if summary:
        _section_header(doc, "Professional Summary")
        sp = doc.add_paragraph()
        _para_spacing(sp, before=3, after=4)
        sr = sp.add_run(summary)
        _set_font(sr, 10.5, color=BLACK)

    # ── WORK EXPERIENCE ───────────────────────────────────────
    experience = data.get("experience") or []
    if experience:
        _section_header(doc, "Work Experience")
        for job in experience:
            # Title line: Job Title (bold) — Company | Date range (right-aligned)
            job_para = doc.add_paragraph()
            _para_spacing(job_para, before=6, after=1)
            _add_tab_stop(job_para, 6.0, "right")

            title_run = job_para.add_run(job.get("title", ""))
            _set_font(title_run, 11, bold=True, color=NAVY)

            sep_run = job_para.add_run("  —  ")
            _set_font(sep_run, 11, color=GREY)

            company_run = job_para.add_run(job.get("company", ""))
            _set_font(company_run, 11, bold=False, color=BLACK)

            # Right-aligned date
            date_str = f"{job.get('start_date', '')} – {job.get('end_date', 'Present')}"
            tab_run = job_para.add_run(f"\t{date_str}")
            _set_font(tab_run, 10, color=LGREY)

            # Location (if present)
            loc = job.get("location", "")
            if loc:
                loc_para = doc.add_paragraph()
                _para_spacing(loc_para, before=0, after=1)
                loc_run = loc_para.add_run(loc)
                _set_font(loc_run, 9.5, italic=True, color=LGREY)

            # Bullets
            for bullet in (job.get("bullets") or []):
                if bullet.strip():
                    _bullet_para(doc, bullet.strip())

    # ── EDUCATION ─────────────────────────────────────────────
    education = data.get("education") or []
    if education:
        _section_header(doc, "Education")
        for edu in education:
            ep = doc.add_paragraph()
            _para_spacing(ep, before=4, after=1)
            _add_tab_stop(ep, 6.0, "right")
            deg_run = ep.add_run(edu.get("degree", ""))
            _set_font(deg_run, 11, bold=True, color=BLACK)
            inst_run = ep.add_run(f"  —  {edu.get('institution', '')}")
            _set_font(inst_run, 11, color=BLACK)
            yr = edu.get("year", "")
            if yr:
                yr_run = ep.add_run(f"\t{yr}")
                _set_font(yr_run, 10, color=LGREY)

    # ── CERTIFICATIONS ────────────────────────────────────────
    certs = data.get("certifications") or []
    if certs:
        _section_header(doc, "Certifications")
        for cert in certs:
            if isinstance(cert, dict):
                text = cert.get("name", "")
                if cert.get("issuer"):
                    text += f" — {cert['issuer']}"
                if cert.get("year"):
                    text += f" ({cert['year']})"
            else:
                text = str(cert)
            if text.strip():
                _bullet_para(doc, text.strip())

    # ── SKILLS ────────────────────────────────────────────────
    skills = data.get("skills") or {}
    skill_lines = []
    if skills.get("core"):
        skill_lines.append(("Core", ", ".join(skills["core"])))
    if skills.get("technical"):
        skill_lines.append(("Technical", ", ".join(skills["technical"])))
    if skills.get("languages"):
        skill_lines.append(("Languages", ", ".join(skills["languages"])))
    if skill_lines:
        _section_header(doc, "Skills")
        for label, content in skill_lines:
            sp = doc.add_paragraph()
            _para_spacing(sp, before=2, after=2)
            lbl = sp.add_run(f"{label}: ")
            _set_font(lbl, 10.5, bold=True, color=NAVY)
            val = sp.add_run(content)
            _set_font(val, 10.5, color=BLACK)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Cover Letter builder ─────────────────────────────────────
def build_cover_letter(data: dict, candidate: dict) -> bytes:
    """
    data: cover letter JSON from AI {recipient, para1, para2, para3, closing}.
    candidate: user dict with name, phone, email fields.
    Returns DOCX bytes.
    """
    from datetime import date
    doc = Document()
    _set_margins(doc, top=1.0, bottom=1.0, left=1.1, right=1.1)

    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)

    # ── CANDIDATE HEADER ──────────────────────────────────────
    name_para = doc.add_paragraph()
    _para_spacing(name_para, before=0, after=2)
    _add_bottom_border(name_para, color_hex="C9A84C", size=12)
    name_run = name_para.add_run((candidate.get("name") or "").upper())
    _set_font(name_run, 15, bold=True, color=NAVY)

    contact_bits = [x for x in [candidate.get("phone"), candidate.get("email")] if x]
    if contact_bits:
        cp = doc.add_paragraph()
        _para_spacing(cp, before=2, after=8)
        cr = cp.add_run("  |  ".join(contact_bits))
        _set_font(cr, 9.5, color=GREY)

    # ── DATE & RECIPIENT ──────────────────────────────────────
    date_p = doc.add_paragraph()
    _para_spacing(date_p, before=0, after=12)
    date_r = date_p.add_run(date.today().strftime("%-d %B %Y"))
    _set_font(date_r, 10.5, color=GREY)

    recip_p = doc.add_paragraph()
    _para_spacing(recip_p, before=0, after=14)
    recip_r = recip_p.add_run(data.get("recipient", "Dear Hiring Team,"))
    _set_font(recip_r, 11, bold=False, color=BLACK)

    # ── THREE PARAGRAPHS ──────────────────────────────────────
    for key in ("para1", "para2", "para3"):
        text = (data.get(key) or "").strip()
        if not text:
            continue
        pp = doc.add_paragraph()
        _para_spacing(pp, before=0, after=10)
        pp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pr = pp.add_run(text)
        _set_font(pr, 11, color=BLACK)

    # ── CLOSING ───────────────────────────────────────────────
    closing_p = doc.add_paragraph()
    _para_spacing(closing_p, before=6, after=2)
    closing_r = closing_p.add_run(data.get("closing", "Yours sincerely,"))
    _set_font(closing_r, 11, color=BLACK)

    sig_p = doc.add_paragraph()
    _para_spacing(sig_p, before=24, after=0)
    sig_r = sig_p.add_run(candidate.get("name") or "")
    _set_font(sig_r, 11, bold=True, color=NAVY)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
