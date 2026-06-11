"""
services/cv_parser.py — PDF/DOCX CV upload parsing.
Extracts full text + best-guess candidate name. Pure-Python (pypdf, python-docx).
"""
import io
import re

def parse_cv(file_bytes, filename):
    """Returns {"text": str, "name": str|None, "error": str|None}"""
    fn = (filename or "").lower()
    try:
        if fn.endswith(".pdf"):
            text = _parse_pdf(file_bytes)
        elif fn.endswith(".docx"):
            text = _parse_docx(file_bytes)
        elif fn.endswith(".txt"):
            text = file_bytes.decode("utf-8", errors="replace")
        else:
            return {"text": "", "name": None, "error": "Unsupported format — use PDF, DOCX, or TXT"}
    except Exception as e:
        return {"text": "", "name": None, "error": f"Parse failed: {str(e)[:150]}"}
    text = re.sub(r"\n{3,}", "\n\n", text).strip()
    if len(text) < 50:
        return {"text": text, "name": None, "error": "File parsed but contains almost no text (scanned image PDF?)"}
    return {"text": text[:15000], "name": _guess_name(text), "error": None}

def _parse_pdf(b):
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(b))
    return "\n".join((p.extract_text() or "") for p in reader.pages)

def _parse_docx(b):
    import docx
    d = docx.Document(io.BytesIO(b))
    parts = [p.text for p in d.paragraphs]
    for table in d.tables:
        for row in table.rows:
            parts.extend(c.text for c in row.cells)
    return "\n".join(parts)

def _guess_name(text):
    """First plausible name line: 2-4 capitalized words, no digits/keywords."""
    bad = re.compile(r"\d|@|curriculum|resume|cv|profile|summary|objective|email|phone|address", re.I)
    for line in text.split("\n")[:10]:
        line = line.strip()
        if not line or bad.search(line):
            continue
        words = line.split()
        if 2 <= len(words) <= 4 and all(w[0].isupper() for w in words if w[0].isalpha()):
            return line[:80]
    return None
